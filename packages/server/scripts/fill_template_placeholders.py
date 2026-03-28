from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "assets" / "templates"
MANIFEST_DIR = TEMPLATE_DIR / "manifests"

SKIP_LABELS = {
    "说明：",
    "当事人信息",
    "诉讼请求",
    "答辩意见",
    "事实与理由",
    "对纠纷解决方式的意愿",
}

SUMMARY_ROW_TITLES = {
    "诉讼请求": ("诉讼请求全文", "request_or_response"),
    "答辩意见": ("答辩意见全文", "request_or_response"),
    "事实与理由": ("事实与理由全文", "facts"),
}


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()


def find_block_title(table, row_index: int) -> str:
    for current in range(row_index - 1, -1, -1):
        texts = [normalize_text(cell.text) for cell in table.rows[current].cells]
        texts = [text for text in texts if text]
        if not texts:
            continue
        unique = list(dict.fromkeys(texts))
        if len(unique) == 1:
            return unique[0]
    return "补充字段"


def guess_source_section(label: str, block_title: str) -> str:
    joined = f"{label} {block_title}"
    if any(token in joined for token in ["原告", "被告", "答辩人", "被答辩人", "第三人", "代理人"]):
        return "parties"
    if any(token in joined for token in ["诉讼请求", "答辩意见", "答辩请求", "本金", "利息", "标的总额", "请求"]):
        return "request_or_response"
    if any(token in joined for token in ["具状人", "答辩人（签字、盖章）", "日期"]):
        return "footer"
    return "facts"


def guess_required(label: str) -> bool:
    if any(token in label for token in ["其他", "证据清单", "诉前保全", "调解", "标的总额"]):
        return False
    return True


def row_has_placeholder(cells) -> bool:
    return any("{{" in cell.text and "}}" in cell.text for cell in cells)


def unique_cells(cells):
    seen = set()
    ordered = []
    for cell in cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        ordered.append(cell)
    return ordered


def ensure_summary_placeholder(document: Document, manifest: dict) -> int:
    used_keys = [field["key"] for field in manifest["fields"]]
    max_index = max((int(key.split("_")[1]) for key in used_keys if key.startswith("field_")), default=0)
    existing_labels = {field["label"] for field in manifest["fields"]}
    added = 0

    for table in document.tables:
        previous_block_title = ""

        for row in table.rows:
            cells = unique_cells(row.cells)
            if not cells:
                continue

            normalized_values = [normalize_text(cell.text) for cell in cells]
            non_empty_values = [value for value in normalized_values if value]
            if not non_empty_values:
                continue

            if len(set(non_empty_values)) == 1 and non_empty_values[0] in SUMMARY_ROW_TITLES:
                previous_block_title = non_empty_values[0]
                continue

            if len(set(non_empty_values)) == 1 and non_empty_values[0].startswith("（可完整表述") and previous_block_title in SUMMARY_ROW_TITLES:
                if row_has_placeholder(cells):
                    continue

                max_index += 1
                key = f"field_{max_index:03d}"
                label, source_section = SUMMARY_ROW_TITLES[previous_block_title]
                cells[-1].add_paragraph(f"{{{{{key}}}}}")
                if label not in existing_labels:
                    manifest["fields"].append({
                        "key": key,
                        "label": label,
                        "blockTitle": previous_block_title,
                        "sourceSection": source_section,
                        "required": False,
                        "multiline": True,
                        "hint": "填写模板说明文字下方的大块概述区域，优先保留原文对应段落。",
                    })
                    existing_labels.add(label)
                    added += 1

    return added


def main() -> None:
    for manifest_path in sorted(MANIFEST_DIR.glob("*.json")):
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        template_path = TEMPLATE_DIR / manifest["templateFile"]
        document = Document(str(template_path))

        used_keys = [field["key"] for field in manifest["fields"]]
        max_index = max(int(key.split("_")[1]) for key in used_keys)
        added = 0

        added += ensure_summary_placeholder(document, manifest)
        used_keys = [field["key"] for field in manifest["fields"]]
        max_index = max(int(key.split("_")[1]) for key in used_keys if key.startswith("field_"))

        for table in document.tables:
            for row_index, row in enumerate(table.rows):
                cells = row.cells
                if len(cells) < 2:
                    continue

                label = normalize_text(cells[0].text)
                if not label or label in SKIP_LABELS or label.startswith("说明"):
                    continue

                right_cells = cells[1:]
                if any("{{field_" in cell.text for cell in right_cells):
                    continue

                empty_cells = [cell for cell in right_cells if not normalize_text(cell.text)]
                if not empty_cells:
                    continue

                block_title = find_block_title(table, row_index)
                max_index += 1
                key = f"field_{max_index:03d}"
                empty_cells[0].text = f"{{{{{key}}}}}"

                manifest["fields"].append({
                    "key": key,
                    "label": label,
                    "blockTitle": block_title,
                    "sourceSection": guess_source_section(label, block_title),
                    "required": guess_required(label),
                    "multiline": True,
                })
                added += 1

        if added > 0:
            document.save(str(template_path))
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            print(f"{manifest_path.name}: added {added} placeholders")
        else:
            print(f"{manifest_path.name}: no changes")


if __name__ == "__main__":
    main()
