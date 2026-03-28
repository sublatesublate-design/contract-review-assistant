from __future__ import annotations

import copy
import json
import re
import shutil
import sys
import zipfile
from collections import OrderedDict, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "assets" / "templates"
MANIFEST_DIR = TEMPLATE_DIR / "manifests"

DEFAULT_SOURCE_DOCX = Path(
    r"C:\Users\24812\Downloads\最高人民法院、司法部、中华全国律师协会关于印发部分案件起诉状答辩状示范文本的通知（法〔2025〕82号）.docx"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
QN = lambda tag: f"{{{W_NS}}}{tag}"

PRIMARY_TITLE_SET = {
    "刑事（附带民事）自诉状",
    "刑事（附带民事）自诉答辩状",
    "民事起诉状",
    "民事答辩状",
    "行政起诉状",
    "行政答辩状",
    "国家赔偿申请书",
    "国家赔偿答辩状",
    "强制执行申请书",
    "执行异议申请书",
    "执行复议申请书",
    "执行监督申请书",
    "不予执行申请书",
    "执行担保申请书",
    "参与分配申请书",
    "确认优先购买权申请书",
    "第三人意见陈述书",
    "暂时解除乘坐飞机、高铁限制措施申请书",
}

AUXILIARY_TITLE_SET = {
    "原告证据清单",
    "被告证据清单",
    "第三人证据清单",
    "原告证据目录",
    "被告证据目录",
    "第三人证据目录",
    "对原告证据的质证意见",
    "对被告证据的质证意见",
    "技术特征对比分析表",
    "现有技术对比分析表",
    "关联行政确权程序及有关行政诉讼案件信息表",
    "同族专利信息表",
    "专利申请有效性分析表",
    "专利有效性分析表",
}

TITLE_SET = PRIMARY_TITLE_SET | AUXILIARY_TITLE_SET

REQUEST_AUXILIARY_TITLE_SET = {
    "原告证据清单",
    "原告证据目录",
}

RESPONSE_AUXILIARY_TITLE_SET = {
    "被告证据清单",
    "第三人证据清单",
    "被告证据目录",
    "第三人证据目录",
    "对原告证据的质证意见",
    "对被告证据的质证意见",
}

CONTEXT_AUXILIARY_TITLE_SET = AUXILIARY_TITLE_SET - REQUEST_AUXILIARY_TITLE_SET - RESPONSE_AUXILIARY_TITLE_SET

CATEGORY_DEFINITIONS: OrderedDict[str, str] = OrderedDict(
    [
        ("criminal_private", "刑事（自诉）"),
        ("civil", "民事"),
        ("commercial", "商事"),
        ("intellectual_property", "知识产权"),
        ("maritime", "海事"),
        ("administrative", "行政"),
        ("environment_resources", "环境资源"),
        ("state_compensation", "国家赔偿"),
        ("enforcement", "执行"),
    ]
)

ENFORCEMENT_TITLES = {
    "强制执行申请书",
    "执行异议申请书",
    "执行复议申请书",
    "执行监督申请书",
    "不予执行申请书",
    "执行担保申请书",
    "参与分配申请书",
    "确认优先购买权申请书",
    "暂时解除乘坐飞机、高铁限制措施申请书",
}

IP_PATTERNS = [
    "著作权",
    "邻接权",
    "商标",
    "专利",
    "植物新品种",
    "商业秘密",
    "技术合同",
    "不正当竞争",
    "垄断",
]

MARITIME_PATTERNS = [
    "船舶",
    "海上",
    "通海水域",
    "船员",
]

ENVIRONMENT_PATTERNS = [
    "环境污染",
    "生态破坏",
    "生态环境",
]

CIVIL_CASES = {
    "离婚纠纷",
    "房屋买卖合同纠纷",
    "房屋租赁合同纠纷",
    "物业服务合同纠纷",
    "劳动争议纠纷",
    "机动车交通事故责任纠纷",
}

COMMERCIAL_CASES = {
    "买卖合同纠纷",
    "金融借款合同纠纷",
    "民间借贷纠纷",
    "信用卡纠纷",
    "融资租赁合同纠纷",
    "建设工程施工合同纠纷",
    "证券虚假陈述责任纠纷",
    "财产损失保险合同纠纷",
    "责任保险合同纠纷",
    "保证保险合同纠纷",
    "人身保险合同纠纷",
}

REQUEST_SECTION_TITLES = {
    "诉讼请求",
    "请求事项",
    "申请事项",
    "申请请求",
    "赔偿请求",
    "执行请求",
    "异议请求",
}

RESPONSE_SECTION_TITLES = {
    "答辩意见",
    "答辩请求",
    "答辩事项",
    "答辩理由",
    "陈述意见",
}

FACT_SECTION_TITLES = {
    "事实与理由",
    "事实和理由",
    "申请理由",
    "异议理由",
    "主要事实与理由",
    "理由",
}

KNOWN_BLOCK_TITLES = (
    {"当事人信息", "落款"}
    | REQUEST_SECTION_TITLES
    | RESPONSE_SECTION_TITLES
    | FACT_SECTION_TITLES
)

SKIP_ROW_LABELS = {
    "说明",
    "当事人信息",
    "诉讼请求",
    "答辩意见",
    "答辩请求",
    "事实与理由",
    "事实和理由",
    "请求事项",
    "申请事项",
    "申请请求",
    "赔偿请求",
    "执行请求",
    "异议请求",
}

FOOTER_LABEL_PATTERNS = [
    "具状人（签字、盖章）",
    "答辩人（签字、盖章）",
    "自诉人（签字、盖章）",
    "赔偿请求人（签字、盖章）",
    "申请人（签字、盖章）",
    "申请执行人（签字、盖章）",
    "第三人（签字、盖章）",
    "日期",
]

BAD_SECTION_PATTERN = re.compile(
    r'<w:pgSz\b[^>]*w:w="11650"[^>]*w:h="16820".*?<w:pgMar\b[^>]*w:top="1"[^>]*w:right="0"[^>]*w:bottom="1"[^>]*w:left="0"',
    re.S,
)


@dataclass
class TemplateSegment:
    title: str
    subtitle: str
    category_id: str
    category_label: str
    orientation: str
    template_id: str
    template_file: str
    manifest_file: str
    file_name_prefix: str
    start_index: int
    end_index: int

    @property
    def label(self) -> str:
        return f"{self.title}{self.subtitle}"

    @property
    def case_title(self) -> str:
        text = self.subtitle.strip()
        if text.startswith("（") and text.endswith("）"):
            return text[1:-1].strip()
        return text


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()


def paragraph_text(element: etree._Element) -> str:
    texts = element.xpath(".//w:t/text()", namespaces=NS)
    return "".join(texts).strip()


def is_paragraph(element: etree._Element) -> bool:
    return element.tag == QN("p")


def is_table(element: etree._Element) -> bool:
    return element.tag == QN("tbl")


def has_visible_text(element: etree._Element) -> bool:
    return bool(normalize_text(paragraph_text(element)))


def strip_paragraph_section_properties(paragraph: etree._Element) -> etree._Element:
    paragraph = copy.deepcopy(paragraph)
    for sect_pr in paragraph.xpath(".//w:sectPr", namespaces=NS):
        parent = sect_pr.getparent()
        if parent is not None:
            parent.remove(sect_pr)
    return paragraph


def choose_section_properties(candidates: Iterable[etree._Element], fallback: etree._Element | None) -> etree._Element | None:
    copied = [copy.deepcopy(candidate) for candidate in candidates if candidate is not None]
    for candidate in copied:
        xml = etree.tostring(candidate, encoding="unicode")
        if not BAD_SECTION_PATTERN.search(xml):
            return candidate
    if copied:
        return copied[-1]
    return copy.deepcopy(fallback) if fallback is not None else None


def merge_tables(tables: list[etree._Element]) -> etree._Element | None:
    if not tables:
        return None
    if len(tables) == 1:
        return copy.deepcopy(tables[0])

    merged = etree.Element(QN("tbl"))
    first_table = tables[0]
    tbl_pr = first_table.find("./w:tblPr", namespaces=NS)
    if tbl_pr is not None:
        merged.append(copy.deepcopy(tbl_pr))

    grids = [table.find("./w:tblGrid", namespaces=NS) for table in tables]
    best_grid = max(
        (grid for grid in grids if grid is not None),
        key=lambda grid: len(grid.findall("./w:gridCol", namespaces=NS)),
        default=None,
    )
    if best_grid is not None:
        merged.append(copy.deepcopy(best_grid))

    for table in tables:
        for row in table.findall("./w:tr", namespaces=NS):
            merged.append(copy.deepcopy(row))

    return merged


def load_document_root(source_docx: Path) -> tuple[bytes, etree._Element, etree._Element]:
    source_bytes = source_docx.read_bytes()
    with zipfile.ZipFile(source_docx) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    body = root.find("./w:body", namespaces=NS)
    if body is None:
        raise RuntimeError("Official DOCX is missing word/document.xml body.")
    return source_bytes, root, body


def previous_non_empty_paragraph_text(children: list[etree._Element], index: int) -> str:
    for current in range(index - 1, -1, -1):
        child = children[current]
        if not is_paragraph(child):
            continue
        text = paragraph_text(child).strip()
        if text:
            return text
    return ""


def next_non_empty_paragraph_text(children: list[etree._Element], index: int) -> str:
    for current in range(index + 1, len(children)):
        child = children[current]
        if not is_paragraph(child):
            continue
        text = paragraph_text(child).strip()
        if text:
            return text
    return ""


def normalize_case_subtitle(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return ""
    if normalized.startswith("（") and normalized.endswith("）"):
        return normalized
    if normalized.startswith("(") and normalized.endswith(")"):
        return f"（{normalized[1:-1].strip()}）"
    return ""


def match_known_title(text: str) -> str | None:
    normalized = normalize_text(text)
    if not normalized:
        return None

    for candidate in TITLE_SET:
        if normalized == candidate:
            return candidate
        if normalized.startswith(f"{candidate}（") or normalized.startswith(f"{candidate} （"):
            return candidate
    return None


def extract_inline_subtitle(text: str, title: str) -> str:
    normalized = normalize_text(text)
    if not normalized.startswith(title):
        return ""
    suffix = normalized[len(title):].strip()
    return normalize_case_subtitle(suffix)


def detect_category(document_title: str, case_title: str) -> str:
    flat_case = case_title.strip()

    if document_title.startswith("刑事"):
        return "criminal_private"
    if document_title.startswith("国家赔偿"):
        return "state_compensation"
    if document_title in ENFORCEMENT_TITLES:
        return "enforcement"
    if any(pattern in flat_case for pattern in ENVIRONMENT_PATTERNS):
        return "environment_resources"
    if any(pattern in flat_case for pattern in MARITIME_PATTERNS):
        return "maritime"
    if any(pattern in flat_case for pattern in IP_PATTERNS) or document_title == "第三人意见陈述书":
        return "intellectual_property"
    if document_title.startswith("行政"):
        return "administrative"
    if flat_case in CIVIL_CASES:
        return "civil"
    if flat_case in COMMERCIAL_CASES:
        return "commercial"
    return "civil"


def detect_orientation(document_title: str) -> str:
    if document_title in REQUEST_AUXILIARY_TITLE_SET:
        return "request"
    if document_title in RESPONSE_AUXILIARY_TITLE_SET:
        return "response"
    if "答辩状" in document_title or document_title == "第三人意见陈述书":
        return "response"
    return "request"


def detect_document_kind(document_title: str) -> str:
    if document_title == "第三人意见陈述书":
        return "third_party_statement"
    if "证据清单" in document_title or "证据目录" in document_title:
        return "evidence_list"
    if "质证意见" in document_title:
        return "cross_examination"
    if "分析表" in document_title:
        return "analysis_table"
    if "信息表" in document_title:
        return "info_table"
    return "main_pleading"


def build_segments(body: etree._Element) -> list[TemplateSegment]:
    children = list(body)
    starts: list[tuple[int, str, str]] = []

    for index, child in enumerate(children):
        if not is_paragraph(child):
            continue
        paragraph_value = normalize_text(paragraph_text(child))
        title = match_known_title(paragraph_value)
        if title is None or title not in PRIMARY_TITLE_SET:
            continue
        if normalize_text(previous_non_empty_paragraph_text(children, index)) == "实例":
            continue
        subtitle = extract_inline_subtitle(paragraph_value, title)
        if not subtitle:
            subtitle = normalize_case_subtitle(next_non_empty_paragraph_text(children, index))
        starts.append((index, title, subtitle))

    category_counters: dict[str, int] = defaultdict(int)
    segments: list[TemplateSegment] = []

    for _position, (start_index, title, subtitle) in enumerate(starts):
        end_index = len(children)
        for next_index in range(start_index + 1, len(children)):
            candidate = children[next_index]
            if not is_paragraph(candidate):
                continue
            candidate_text = normalize_text(paragraph_text(candidate))
            if candidate_text == "实例":
                end_index = next_index
                break
            if match_known_title(candidate_text) and normalize_text(previous_non_empty_paragraph_text(children, next_index)) != "实例":
                end_index = next_index
                break

        orientation = detect_orientation(title)
        case_title = subtitle[1:-1] if subtitle.startswith("（") and subtitle.endswith("）") else subtitle
        category_id = detect_category(title, case_title)
        category_label = CATEGORY_DEFINITIONS[category_id]
        category_counters[category_id] += 1
        sequence = category_counters[category_id]
        template_id = f"{category_id}_{sequence:03d}"
        segments.append(
            TemplateSegment(
                title=title,
                subtitle=subtitle,
                category_id=category_id,
                category_label=category_label,
                orientation=orientation,
                template_id=template_id,
                template_file=f"{template_id}.docx",
                manifest_file=f"{template_id}.json",
                file_name_prefix=f"{title}{subtitle}",
                start_index=start_index,
                end_index=end_index,
            )
        )

    return segments


def build_segment_document_xml(source_root: etree._Element, body: etree._Element, segment: TemplateSegment) -> bytes:
    selected = [copy.deepcopy(child) for child in list(body)[segment.start_index:segment.end_index]]
    fallback_sect = body.find("./w:sectPr", namespaces=NS)
    section_candidates = body.xpath(".//w:sectPr", namespaces=NS)
    section = choose_section_properties(section_candidates, fallback_sect)

    new_children: list[etree._Element] = []
    tables_to_merge: list[etree._Element] = []
    first_table_inserted = False

    def flush_tables() -> None:
        nonlocal first_table_inserted
        if tables_to_merge and not first_table_inserted:
            merged = merge_tables(tables_to_merge)
            if merged is not None:
                new_children.append(merged)
            first_table_inserted = True
        tables_to_merge.clear()

    for child in selected:
        if is_table(child):
            tables_to_merge.append(child)
            continue

        if is_paragraph(child):
            paragraph = strip_paragraph_section_properties(child)
            if tables_to_merge and not has_visible_text(paragraph):
                continue
            flush_tables()
            if not has_visible_text(paragraph):
                if not new_children:
                    continue
                if new_children and is_table(new_children[-1]):
                    continue
            new_children.append(paragraph)

    flush_tables()

    while new_children and is_paragraph(new_children[-1]) and not has_visible_text(new_children[-1]):
        new_children.pop()

    working_root = copy.deepcopy(source_root)
    working_body = working_root.find("./w:body", namespaces=NS)
    if working_body is None:
        raise RuntimeError("Generated template is missing body.")

    for child in list(working_body):
        working_body.remove(child)

    for child in new_children:
        working_body.append(child)
    if section is not None:
        working_body.append(section)

    return etree.tostring(working_root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def write_template_docx(source_docx: Path, source_bytes: bytes, document_xml: bytes, output_path: Path) -> None:
    with zipfile.ZipFile(source_docx) as source_archive:
        with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as target_archive:
            for info in source_archive.infolist():
                data = source_archive.read(info.filename)
                if info.filename == "word/document.xml":
                    data = document_xml
                target_archive.writestr(info, data)


def unique_row_cells(row) -> list:
    seen = set()
    ordered = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        ordered.append(cell)
    return ordered


def insert_placeholder_into_cell(cell, token: str) -> None:
    if cell.paragraphs:
        last_text = normalize_text(cell.paragraphs[-1].text)
        if not last_text:
            cell.paragraphs[-1].text = token
            return
    cell.add_paragraph(token)


def insert_placeholder_into_paragraph(paragraph, token: str) -> None:
    if "{{" in paragraph.text and "}}" in paragraph.text:
        return
    paragraph.add_run(token)


def find_block_title(table, row_index: int) -> str:
    for current in range(row_index - 1, -1, -1):
        cells = unique_row_cells(table.rows[current])
        texts = [normalize_text(cell.text) for cell in cells if normalize_text(cell.text)]
        if not texts:
            continue
        unique_texts = list(dict.fromkeys(texts))
        if len(unique_texts) == 1:
            matched = match_block_title(unique_texts[0])
            if matched:
                return matched
    return "补充字段"


def guess_source_section(label: str, block_title: str, orientation: str) -> str:
    joined = f"{label} {block_title}"

    if any(token in joined for token in ["原告", "被告", "答辩人", "被答辩人", "第三人", "申请人", "被申请人", "自诉人", "赔偿请求人", "赔偿义务机关", "申请执行人", "被执行人", "异议人", "案外人", "利害关系人", "代理人"]):
        return "parties"

    if any(token in joined for token in ["具状人", "答辩人（签字、盖章）", "申请人（签字、盖章）", "第三人（签字、盖章）", "日期"]):
        return "footer"

    if block_title in REQUEST_SECTION_TITLES or block_title in RESPONSE_SECTION_TITLES:
        return "request_or_response"

    if any(token in joined for token in ["诉讼请求", "答辩意见", "答辩请求", "请求事项", "申请事项", "赔偿请求", "执行请求"]):
        return "request_or_response"

    if any(token in joined for token in ["事实与理由", "事实和理由", "申请理由", "异议理由"]):
        return "facts"

    if orientation == "response" and "陈述意见" in joined:
        return "request_or_response"

    return "facts"


def guess_required(label: str, block_title: str) -> bool:
    text = f"{label} {block_title}"
    if any(token in text for token in ["其他", "证据清单", "可另附页", "先行调解", "调解", "了解", "考虑", "标的总额", "诉前保全", "签字", "盖章", "日期"]):
        return False
    return True


def match_block_title(text: str) -> str | None:
    normalized = normalize_text(text)
    if not normalized:
        return None
    if (
        normalized.startswith("说明")
        or normalized.startswith("（")
        or "特别提示" in normalized
        or "相关内容请在下方要素式表格中填写" in normalized
    ):
        return None
    for title in sorted(KNOWN_BLOCK_TITLES, key=len, reverse=True):
        if normalized == title or normalized.startswith(title):
            return title
    if len(normalized) <= 40:
        return normalized
    return None


def is_summary_instruction_text(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized.startswith("（"):
        return False
    return (
        normalized.startswith("（可完整表述")
        or normalized.startswith("（请概况描述")
        or "相关内容请在下方要素式表格中填写" in normalized
        or "相关具体内容请在下方要素式表格中填写" in normalized
    )


def build_manifest_for_template(segment: TemplateSegment, template_path: Path) -> dict:
    document = Document(str(template_path))
    fields: list[dict] = []
    field_index = 1
    existing_labels: set[str] = set()

    def add_field(
        label: str,
        block_title: str,
        source_section: str,
        multiline: bool,
        hint: str | None = None,
        required: bool | None = None,
    ) -> str:
        nonlocal field_index
        key = f"field_{field_index:03d}"
        field_index += 1
        normalized_hint = hint if hint else None
        field = {
            "key": key,
            "label": label,
            "blockTitle": block_title,
            "sourceSection": source_section,
            "required": guess_required(label, block_title) if required is None else required,
            "multiline": multiline,
        }
        if normalized_hint:
            field["hint"] = normalized_hint
        fields.append(field)
        existing_labels.add(label)
        return key

    for table in document.tables:
        current_block_title = ""
        for row_index, row in enumerate(table.rows):
            cells = unique_row_cells(row)
            if not cells:
                continue

            normalized_values = [normalize_text(cell.text) for cell in cells]
            non_empty_values = [value for value in normalized_values if value]
            if not non_empty_values:
                continue

            unique_texts = list(dict.fromkeys(non_empty_values))
            if len(unique_texts) == 1:
                candidate_title = unique_texts[0]
                matched_block_title = match_block_title(candidate_title)
                if matched_block_title:
                    current_block_title = matched_block_title
                    continue

                if is_summary_instruction_text(candidate_title) and current_block_title:
                    summary_label = f"{current_block_title}全文"
                    if summary_label not in existing_labels:
                        token = add_field(
                            label=summary_label,
                            block_title=current_block_title,
                            source_section="request_or_response" if current_block_title in REQUEST_SECTION_TITLES | RESPONSE_SECTION_TITLES else "facts",
                            multiline=True,
                            hint="填写模板说明文字下方的大段概述区域，优先保留原文对应段落。",
                            required=False,
                        )
                        insert_placeholder_into_cell(cells[-1], f"{{{{{token}}}}}")
                    continue

            if len(cells) < 2:
                continue

            label = normalize_text(cells[0].text)
            if (
                not label
                or label in SKIP_ROW_LABELS
                or label.startswith("（")
                or label.startswith("说明")
                or "特别提示" in label
            ):
                continue

            target_cell = cells[-1]
            hint = normalize_text(target_cell.text)
            if "{{" in target_cell.text and "}}" in target_cell.text:
                continue

            if not current_block_title:
                current_block_title = find_block_title(table, row_index)

            token = add_field(
                label=label,
                block_title=current_block_title,
                source_section=guess_source_section(label, current_block_title, segment.orientation),
                multiline=True,
                hint=hint or None,
            )
            insert_placeholder_into_cell(target_cell, f"{{{{{token}}}}}")

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text or "{{" in text:
            continue
        for footer_label in FOOTER_LABEL_PATTERNS:
            if footer_label in text and footer_label not in existing_labels:
                token = add_field(
                    label=footer_label,
                    block_title="落款",
                    source_section="footer",
                    multiline=False,
                    required=False,
                )
                insert_placeholder_into_paragraph(paragraph, f"{{{{{token}}}}}")
                break

    document.save(str(template_path))

    return {
        "templateId": segment.template_id,
        "categoryId": segment.category_id,
        "categoryLabel": segment.category_label,
        "orientation": segment.orientation,
        "documentKind": detect_document_kind(segment.title),
        "documentTitle": segment.title,
        "caseTitle": segment.case_title,
        "label": segment.label,
        "templateFile": segment.template_file,
        "manifestFile": segment.manifest_file,
        "fileNamePrefix": segment.file_name_prefix,
        "fields": fields,
    }


def build_catalog(manifests: list[dict]) -> list[dict]:
    groups: OrderedDict[str, dict] = OrderedDict(
        (
            category_id,
            {
                "id": category_id,
                "label": label,
                "items": [],
            },
        )
        for category_id, label in CATEGORY_DEFINITIONS.items()
    )

    for manifest in manifests:
        groups[manifest["categoryId"]]["items"].append(
            {
                "templateId": manifest["templateId"],
                "categoryId": manifest["categoryId"],
                "categoryLabel": manifest["categoryLabel"],
                "documentKind": manifest["documentKind"],
                "documentTitle": manifest["documentTitle"],
                "caseTitle": manifest["caseTitle"],
                "label": manifest["label"],
                "orientation": manifest["orientation"],
                "templateFile": manifest["templateFile"],
                "manifestFile": manifest["manifestFile"],
                "fileNamePrefix": manifest["fileNamePrefix"],
            }
        )

    return [group for group in groups.values() if group["items"]]


def ensure_clean_output_dirs() -> None:
    if TEMPLATE_DIR.exists():
        for child in TEMPLATE_DIR.iterdir():
            if child.is_dir():
                shutil.rmtree(child)
            else:
                child.unlink()
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    MANIFEST_DIR.mkdir(parents=True, exist_ok=True)


def main() -> None:
    source_docx = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE_DOCX
    if not source_docx.exists():
        raise FileNotFoundError(f"Official source DOCX not found: {source_docx}")

    ensure_clean_output_dirs()
    source_bytes, source_root, body = load_document_root(source_docx)
    segments = build_segments(body)

    manifests: list[dict] = []
    for segment in segments:
        document_xml = build_segment_document_xml(source_root, body, segment)
        template_path = TEMPLATE_DIR / segment.template_file
        write_template_docx(source_docx, source_bytes, document_xml, template_path)
        manifest = build_manifest_for_template(segment, template_path)
        manifest_path = MANIFEST_DIR / segment.manifest_file
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        manifests.append(manifest)

    catalog = build_catalog(manifests)
    (TEMPLATE_DIR / "catalog.json").write_text(json.dumps(catalog, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    print(f"Generated {len(manifests)} official templates under {TEMPLATE_DIR}")
    for category in catalog:
        print(f"- {category['label']}: {len(category['items'])}")


if __name__ == "__main__":
    main()
